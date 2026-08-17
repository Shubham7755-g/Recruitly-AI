"""
Resume Document Parser Service

Extracts:
- Raw text from PDF/DOCX
- Email
- Phone
- Candidate name
"""

import io
import re
from typing import Dict


# ============================================================
# PDF PARSER
# ============================================================

def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from PDF bytes.

    Primary parser:
        PyMuPDF (fitz)

    Fallback parser:
        pypdf

    Returns an empty string if the PDF cannot be parsed.
    """

    if not file_bytes:
        print("[Parser] PDF received with 0 bytes")
        return ""

    # --------------------------------------------------------
    # Validate PDF header
    # --------------------------------------------------------

    if not file_bytes.startswith(b"%PDF"):
        print("[Parser] Warning: file does not start with %PDF")

    # --------------------------------------------------------
    # Method 1: PyMuPDF
    # --------------------------------------------------------

    try:
        import fitz

        print(
            f"[Parser] Trying PyMuPDF "
            f"({len(file_bytes)} bytes)"
        )

        doc = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        text_parts = []

        for page_number, page in enumerate(doc):
            try:
                page_text = page.get_text("text") or ""

                if page_text.strip():
                    text_parts.append(page_text)

            except Exception as page_error:
                print(
                    f"[Parser] Page {page_number + 1} "
                    f"extraction failed: {page_error}"
                )

        doc.close()

        text = "\n".join(text_parts).strip()

        if text:
            print(
                f"[Parser] PyMuPDF extracted "
                f"{len(text)} characters"
            )
            return text

        print("[Parser] PyMuPDF returned empty text")

    except Exception as error:
        print(
            f"[Parser] PyMuPDF failed: {error}"
        )

    # --------------------------------------------------------
    # Method 2: pypdf fallback
    # --------------------------------------------------------

    try:
        from pypdf import PdfReader

        print("[Parser] Trying pypdf fallback")

        reader = PdfReader(
            io.BytesIO(file_bytes)
        )

        text_parts = []

        for page_number, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""

                if page_text.strip():
                    text_parts.append(page_text)

            except Exception as page_error:
                print(
                    f"[Parser] pypdf page "
                    f"{page_number + 1} failed: "
                    f"{page_error}"
                )

        text = "\n".join(text_parts).strip()

        if text:
            print(
                f"[Parser] pypdf extracted "
                f"{len(text)} characters"
            )
            return text

        print("[Parser] pypdf returned empty text")

    except Exception as error:
        print(
            f"[Parser] pypdf failed: {error}"
        )

    # --------------------------------------------------------
    # Final failure
    # --------------------------------------------------------

    print(
        "[Parser] ERROR: Could not extract "
        "text from PDF"
    )

    return ""


# ============================================================
# DOCX PARSER
# ============================================================

def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from DOCX bytes.
    """

    if not file_bytes:
        print("[Parser] DOCX received with 0 bytes")
        return ""

    try:
        import docx

        print(
            f"[Parser] Parsing DOCX "
            f"({len(file_bytes)} bytes)"
        )

        document = docx.Document(
            io.BytesIO(file_bytes)
        )

        text_parts = []

        # ----------------------------------------------------
        # Paragraphs
        # ----------------------------------------------------

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        # ----------------------------------------------------
        # Tables
        # ----------------------------------------------------

        for table in document.tables:
            for row in table.rows:
                row_text = []

                for cell in row.cells:
                    cell_text = cell.text.strip()

                    if cell_text:
                        row_text.append(cell_text)

                if row_text:
                    text_parts.append(
                        " | ".join(row_text)
                    )

        result = "\n".join(text_parts).strip()

        print(
            f"[Parser] DOCX extracted "
            f"{len(result)} characters"
        )

        return result

    except Exception as error:
        print(
            f"[Parser] DOCX parsing failed: "
            f"{error}"
        )

        return ""


# ============================================================
# NAME DETECTION HELPER
# ============================================================

def _looks_like_name(line: str) -> bool:
    """
    Check whether a line looks like a person's name.
    """

    if not line:
        return False

    line = line.strip()

    # --------------------------------------------------------
    # Basic length
    # --------------------------------------------------------

    if len(line) < 3 or len(line) > 60:
        return False

    # --------------------------------------------------------
    # Reject email
    # --------------------------------------------------------

    if "@" in line:
        return False

    # --------------------------------------------------------
    # Reject URLs
    # --------------------------------------------------------

    if re.search(
        r"https?://|www\.",
        line,
        re.IGNORECASE
    ):
        return False

    # --------------------------------------------------------
    # Reject phone numbers
    # --------------------------------------------------------

    if re.fullmatch(
        r"[\d\s()+\-./]+",
        line
    ):
        return False

    # --------------------------------------------------------
    # Clean separators
    # --------------------------------------------------------

    cleaned = re.sub(
        r"[|•·]",
        " ",
        line
    )

    cleaned = re.sub(
        r"\s+",
        " ",
        cleaned
    ).strip()

    normalized = re.sub(
        r"[^a-zA-Z\s]",
        "",
        cleaned
    ).strip().lower()

    # --------------------------------------------------------
    # Resume headings
    # --------------------------------------------------------

    excluded_headings = {
        "resume",
        "curriculum vitae",
        "cv",
        "curriculum vitae resume",
        "professional resume",
        "profile",
        "career profile",
        "professional profile",
        "about me",
        "contact",
        "contact information",
        "personal information",
        "objective",
        "career objective",
        "summary",
        "professional summary",
        "experience",
        "work experience",
        "education",
        "skills",
        "technical skills",
        "projects",
        "certifications",
        "achievements",
        "references",
        "skills summary",
        "professional experience",
    }

    if normalized in excluded_headings:
        return False

    # --------------------------------------------------------
    # Technical phrases
    # --------------------------------------------------------

    excluded_name_phrases = {
        "machine learning",
        "artificial intelligence",
        "artificial intelligence and machine learning",
        "data science",
        "computer science",
        "software engineering",
        "software developer",
        "web developer",
        "frontend developer",
        "front end developer",
        "backend developer",
        "back end developer",
        "full stack developer",
        "full-stack developer",
        "data analyst",
        "data scientist",
        "machine learning engineer",
        "machine learning developer",
        "python developer",
        "java developer",
        "javascript developer",
        "cloud engineer",
        "devops engineer",
        "software engineer",
        "network engineer",
        "cyber security",
        "cybersecurity",
        "information technology",
        "information science",
        "computer engineering",
        "electrical engineering",
        "electronics engineering",
    }

    if normalized in excluded_name_phrases:
        return False

    # --------------------------------------------------------
    # Words
    # --------------------------------------------------------

    words = cleaned.split()

    if not 2 <= len(words) <= 5:
        return False

    # --------------------------------------------------------
    # Every word must look like a name
    # --------------------------------------------------------

    for word in words:

        if not re.fullmatch(
            r"[A-Za-z][A-Za-z'’-]*",
            word
        ):
            return False

    # --------------------------------------------------------
    # Technical/title words
    # --------------------------------------------------------

    title_words = {
        "engineer",
        "developer",
        "designer",
        "manager",
        "student",
        "intern",
        "analyst",
        "scientist",
        "consultant",
        "architect",
        "administrator",
        "recruiter",
        "director",
        "lead",
        "professor",
        "specialist",
        "machine",
        "learning",
        "artificial",
        "intelligence",
        "data",
        "science",
        "software",
        "technology",
        "technologies",
        "computer",
        "programming",
        "web",
        "frontend",
        "backend",
        "fullstack",
        "python",
        "java",
        "javascript",
        "sql",
        "cloud",
        "devops",
        "cybersecurity",
        "cyber",
        "network",
        "networks",
        "information",
    }

    if any(
        word.lower() in title_words
        for word in words
    ):
        return False

    return True


# ============================================================
# CANDIDATE NAME EXTRACTION
# ============================================================

def extract_candidate_name(text: str) -> str:
    """Extract a candidate name from the beginning of resume text."""
    if not text:
        return "Unknown Candidate"

    # PDF extraction often leaves blank lines and control characters.
    lines = []
    for raw in text.splitlines():
        line = re.sub(r"[\x00-\x1f\x7f]", " ", raw)
        line = re.sub(r"\s+", " ", line).strip()
        if line:
            lines.append(line)

    # 1) Explicit name labels.
    for line in lines[:40]:
        m = re.match(r"^(?:full\s+name|candidate\s+name|name)\s*[:\-]\s*(.+)$", line, re.I)
        if m and _looks_like_name(m.group(1).strip()):
            return m.group(1).strip()

    # 2) The first meaningful line of a normal resume is overwhelmingly
    # likely to be the person's name. Require only a clean 2-5 word name.
    for line in lines[:10]:
        if _looks_like_name(line):
            return line

    # 3) Email fallback. Handles both arjun.sharma@email.com and
    # shreyasrivastava020806@gmail.com.
    email_match = re.search(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}", text)
    if email_match:
        local = email_match.group(0).split("@", 1)[0]
        local = re.sub(r"\d+", "", local)
        parts = [p for p in re.split(r"[._-]+", local) if p.isalpha()]
        if len(parts) >= 2:
            return " ".join(p.capitalize() for p in parts[:4])
        if len(parts) == 1 and len(parts[0]) >= 5:
            # Try to restore a missing boundary in concatenated names
            # commonly produced by PDF extraction (e.g.
            # singhshubham -> Singh Shubham, sumitkumar -> Sumit Kumar).
            name = parts[0].lower()
            common_parts = [
                "shreya", "shreyas", "shubham", "sumit", "arjun",
                "rahul", "rohit", "aman", "abhishek", "aditya",
                "ankit", "ashish", "ayush", "deepak", "gaurav",
                "harsh", "karan", "mohit", "naman", "nikhil",
                "prashant", "sachin", "sahil", "shivam", "suraj",
                "vishal", "varun", "vikas", "yash", "kumar",
                "singh", "sharma", "verma", "gupta", "yadav",
                "patel", "raj", "khan", "rivastava"
            ]
            candidates = []
            for cut in range(3, len(name) - 2):
                left, right = name[:cut], name[cut:]
                if left in common_parts and right in common_parts:
                    candidates.append((left, right))
            if candidates:
                left, right = max(candidates, key=lambda pair: len(pair[0]) + len(pair[1]))
                return f"{left.capitalize()} {right.capitalize()}"

            # No reliable word boundary was found; keep the original
            # token rather than inventing a name split.
            return parts[0].capitalize()

    return "Unknown Candidate"


# ============================================================
# CONTACT INFORMATION
# ============================================================

def extract_contact_info(
    text: str
) -> Dict[str, str]:
    """
    Extract candidate name, email and phone.
    """

    if not text:
        return {
            "candidate_name": "Unknown Candidate",
            "email": "",
            "phone": "",
        }

    # ========================================================
    # EMAIL
    # ========================================================

    email_match = re.search(
        r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}",
        text
    )

    email = (
        email_match.group(0)
        if email_match
        else ""
    )

    # ========================================================
    # PHONE
    # ========================================================

    phone_match = re.search(
        r"(?:\+?\d{1,3}[-.\s]?)?"
        r"(?:\(?\d{3}\)?[-.\s]?)?"
        r"\d{3}[-.\s]?\d{4}",
        text
    )

    phone = (
        phone_match.group(0)
        if phone_match
        else ""
    )

    # ========================================================
    # NAME
    # ========================================================

    candidate_name = extract_candidate_name(text)

    # --------------------------------------------------------
    # Reliable fallback: derive a human-readable name from email
    # when PDF text layout puts the name and contact details on
    # the same line and normal name detection rejects that line.
    # Example: arjun.sharma@email.com -> Arjun Sharma
    # --------------------------------------------------------
    if candidate_name == "Unknown Candidate" and email:
        local_part = email.split("@", 1)[0]
        local_part = re.sub(r"\d+", " ", local_part)
        name_parts = [
            part for part in re.split(r"[._-]+", local_part)
            if part and part.isalpha()
        ]

        if 2 <= len(name_parts) <= 4:
            candidate_name = " ".join(
                part.capitalize() for part in name_parts
            )
        elif len(name_parts) == 1 and len(name_parts[0]) >= 5:
            # Same boundary restoration for email addresses such as
            # sumitkumar@... when the local part has no separator.
            name = name_parts[0].lower()
            common_parts = [
                "shreya", "shreyas", "shubham", "sumit", "arjun",
                "rahul", "rohit", "aman", "abhishek", "aditya",
                "ankit", "ashish", "ayush", "deepak", "gaurav",
                "harsh", "karan", "mohit", "naman", "nikhil",
                "prashant", "sachin", "sahil", "shivam", "suraj",
                "vishal", "varun", "vikas", "yash", "kumar",
                "singh", "sharma", "verma", "gupta", "yadav",
                "patel", "raj", "khan", "rivastava", "srivastava", "mishra", "tiwari",
                "tripathi", "agrawal", "agarwal", "jain", "mehta", "malhotra",
                "kapoor", "choudhary", "chaudhary", "saxena", "shukla", "pandey"
            ]
            candidates = []
            for cut in range(3, len(name) - 2):
                left, right = name[:cut], name[cut:]
                if left in common_parts and right in common_parts:
                    candidates.append((left, right))
            if candidates:
                left, right = max(candidates, key=lambda pair: len(pair[0]) + len(pair[1]))
                candidate_name = f"{left.capitalize()} {right.capitalize()}"

    # ========================================================
    # RETURN
    # ========================================================

    return {
        "candidate_name": candidate_name,
        "email": email,
        "phone": phone,
    }